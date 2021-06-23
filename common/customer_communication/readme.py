from __future__ import annotations

import os
from pathlib import Path
from zipfile import ZipFile

import pandas as pd
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Inches, Pt, RGBColor
from pandas import DataFrame


class ReadmeBuilder:
    def __init__(
        self,
        df: DataFrame,
        folder: ZipFile,
        fname: str,
        codebook: str | None = None,
        documentation: str | None = None,
        coded_input: dict[str, str] | None = None,
        to_zip: bool = True,
    ):

        self.data = df
        self.file_name = fname
        self.folder = folder
        self.to_zip = to_zip
        self.readme = "Readme.docx"
        if not coded_input:
            self.client_name = input("Client Name: ")
            self.objective = input(
                "give a short description of the objective of the data request and list the requirements"
            )
            self.version = input("Version:")
        else:
            self.client_name = coded_input["client_name"]

            self.objective = coded_input["objective"]
            self.version = coded_input["version"]

        self.codebook = codebook
        self.product_doc = documentation

        self.num_records = self.data.shape[0]
        self.num_var = self.data.shape[1]
        self.empty_cells = self.data.isnull().sum().sum()  # noqa
        self.perc_empty_cells = (
            self.empty_cells / (self.num_records * self.num_var) * 100
        ).round(2)
        self.date_fields = len((self.data.select_dtypes(include="datetime")).columns)
        self.bool_fields = (self.data.apply(lambda x: x.nunique()) == 2).sum()
        self.text_fields = len((self.data.select_dtypes(include="object")).columns)
        self.numeric_fields = len((self.data.select_dtypes(include="number")).columns)
        self.file_size = (self.data.memory_usage(deep=True).sum() / 1024 ** 2).round(
            2
        )  # noqa
        self.cols = self.data.columns.to_list()

        self.logo = os.path.join(Path(__file__).parent / "Matrixian_logo.png")

        document = Document()

        # STYLES
        styles = document.styles
        style = styles.add_style("Normal Text", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Iconic Medium"
        style.font.size = Pt(10)

        styles = document.styles
        style = styles.add_style("Header 2", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Iconic Medium"
        style.font.size = Pt(14)
        style.font.color.rgb = RGBColor(3, 121, 96)
        style.font.bold = True

        style = styles.add_style("Header 1", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Iconic Medium"
        style.font.size = Pt(40)
        style.font.color.rgb = RGBColor(3, 121, 96)
        style.font.bold = True

        style = styles.add_style("Header 3", WD_STYLE_TYPE.PARAGRAPH)
        style.font.name = "Iconic Medium"
        style.font.size = Pt(22)
        style.font.color.rgb = RGBColor(3, 121, 96)
        style.font.bold = True

        # HEADER
        header = document.sections[0].header
        p = header.add_paragraph("")
        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r = p.add_run()
        r.add_picture(fr"{self.logo}", width=Inches(1.5))

        # Title
        h = document.add_heading("Readme", 1)
        h.style = document.styles["Header 1"]

        p = document.add_paragraph(
            f'Bestandslevering van {len(self.data)} rijen voor "{self.client_name}"'
        )
        p.style = document.styles["Normal Text"]
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(20)

        p.style = document.styles["Normal Text"]

        h = document.add_heading("Algemene Informatie")
        h.style = document.styles["Header 3"]
        h.paragraph_format.space_after = Pt(10)

        # GETTING STARTED
        h = document.add_heading("INTRODUCTIE")
        h.style = document.styles["Header 2"]
        h.paragraph_format.space_before = Pt(10)
        h.paragraph_format.space_after = Pt(4)
        p = document.add_paragraph(
            "Als begeleiding bij de voor u geprepareerde dataset ontvangt u een automatisch "
            "gegenereerd readme bestand en codeboek.\n\nDit readme document bevat een opsomming "
            "van de uitgeleverde bestanden, een beknopte omschrijving van de dataset en onze "
            "contactinformatie.\n\nHet codeboek bevat een meer gedetailleerde beschrijving van de "
            "dataset. Hierin is het data profiel en de verbose omschrijving van elke kolom in te "
            "zien.\n\nVoor vragen kunt u natuurlijk altijd contact met ons opnemen.\n\n"
        )
        p.style = document.styles["Normal Text"]
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY_LOW

        # MATRIXIAN
        h = document.add_heading("MATRIXIAN GROUP").style = document.styles["Header 2"]
        h.paragraph_format.space_after = Pt(4)
        p = document.add_paragraph(
            "www.matrixiangroup.com\n"
            "info@matrixiangroup.com\n"
            "+31 (0)20 244 0145\n"
            f"Klantnaam: {self.client_name}\n"
        )
        p.style = document.styles["Normal Text"]

        # BESTANDEN
        h = document.add_heading("BESTANDEN")
        h.style = document.styles["Header 2"]
        h.paragraph_format.space_after = Pt(4)
        h.paragraph_format.space_before = Pt(10)
        p = document.add_paragraph(f"- {self.file_name} ", style="List Bullet")
        p.style = document.styles["Normal Text"]
        p = document.add_paragraph(f"- {self.readme} ", style="List Bullet")
        p.style = document.styles["Normal Text"]
        if self.codebook:
            p = document.add_paragraph(f"- {self.codebook} ", style="List Bullet")
            p.style = document.styles["Normal Text"]
        if self.product_doc:
            p = document.add_paragraph(f"- {self.product_doc} ", style="List Bullet")
            p.style = document.styles["Normal Text"]

        # EENHEDEN
        h = document.add_heading("EENHEDEN")
        h.style = document.styles["Header 2"]
        h.paragraph_format.space_after = Pt(6)
        p = document.add_paragraph(
            """Ruimtelijk: metrisch, meters
Datums: JJJJ-MM-DD
Boolean: 1 = True, 0 = False
Valuta: in euro's (EUR/€)"""
        )
        p.style = document.styles["Normal Text"]

        document.add_page_break()
        h = document.add_heading("Dataoverzicht", 1)
        h.style = document.styles["Header 3"]
        h.paragraph_format.space_after = Pt(10)

        # OMSCRHIJVING
        h = document.add_heading("OMSCHRIJVING")
        h.style = document.styles["Header 2"]
        h.paragraph_format.space_after = Pt(4)
        p = document.add_paragraph(
            f"""{self.objective}
"""
        )
        p.style = document.styles["Normal Text"]

        # DATA OVERZICHT
        h = document.add_heading("DATA TYPE")
        h.paragraph_format.space_before = Pt(20)
        h.paragraph_format.space_after = Pt(4)
        h.style = document.styles["Header 2"]
        data = {
            "Omschrijving": [
                "Aantal kolommen",
                "Aantal rijen",
                "Lege velden",
                "Numerieke kolommen",
                "Tekstuele kolommen",
                "Datum kolommen",
                "Boolean kolommen",
            ],
            "Waarde": [
                self.num_var,
                self.num_records,
                f"{self.empty_cells} ({self.perc_empty_cells})%",
                self.numeric_fields,
                self.text_fields,
                self.date_fields,
                self.bool_fields,
            ],
        }

        table1 = pd.DataFrame(data, columns=["Omschrijving", "Waarde"])
        table = document.add_table(table1.shape[0] + 1, table1.shape[1])
        table.style = "Table Grid"
        # add the header rows.
        for j in range(table1.shape[-1]):
            table.cell(0, j).text = table1.columns[j]

        # add the rest of the data frame
        for i in range(table1.shape[0]):
            for j in range(table1.shape[-1]):
                table.cell(i + 1, j).text = str(table1.values[i, j])

        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].style = document.styles["Normal Text"]
                row.height = Cm(0.6)
                cell.width = Cm(1)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # DATA VELDEN
        h = document.add_heading("DATA VELDEN")
        h.style = document.styles["Header 2"]
        h.paragraph_format.space_before = Pt(20)
        h.paragraph_format.space_after = Pt(4)

        table = document.add_table(rows=0, cols=1)
        table.style = "Table Grid"
        table.allow_autofit = False

        for col in self.cols:
            cells = table.add_row().cells
            cells[0].text = col
            cells[0].width = Inches(0.5)

        for row in table.rows:
            row.height = Cm(0.6)
            row.width = Cm(5)
            for cell in row.cells:
                cell.paragraphs[0].style = document.styles["Normal Text"]
                cell.width = Cm(5)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        document.save(self.readme)

    def write_file(self) -> None:

        self.folder.write(self.readme)
        if self.to_zip:
            os.remove(self.readme)


def readme_exe(
    df: DataFrame,
    folder: ZipFile,
    fname: str,
    codebook: str | None = None,
    documentation: str | None = None,
    coded_input: dict[str, str] | None = None,
    to_zip: bool = True,
) -> None:
    rm_m = ReadmeBuilder(
        df=df,
        folder=folder,
        fname=fname,
        codebook=codebook,
        documentation=documentation,
        coded_input=coded_input,
        to_zip=to_zip,
    )
    rm_m.write_file()
